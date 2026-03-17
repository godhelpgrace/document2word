"""
DOCX renderer.

Converts a Document model to an editable .docx file.
Uses page background images with overlaid text boxes for layout preservation.
"""

import io
import logging
from itertools import count
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from model.document import Document, Page, Block, BlockType

logger = logging.getLogger(__name__)

# Points to EMU conversion
PT_TO_EMU = 12700
INCH_TO_EMU = 914400

# Ensure required namespaces are registered
nsmap.setdefault("v", "urn:schemas-microsoft-com:vml")
nsmap.setdefault("o", "urn:schemas-microsoft-com:office:office")
nsmap.setdefault("wp", "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing")
nsmap.setdefault("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
nsmap.setdefault("wps", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")

# Unique id generator for shapes/docPr
_shape_id_counter = count(1)


def render_document_to_docx(document: Document, output_path: str) -> str:
    """
    Render a Document model to a DOCX file.

    Strategy:
    - Each page becomes a section in the DOCX
    - Page dimensions match the original PDF
    - Background image is placed as the page background
    - Text blocks are added as positioned paragraphs

    Args:
        document: The unified Document model
        output_path: Path to save the DOCX file

    Returns:
        The output path
    """
    docx = DocxDocument()

    # Remove default empty paragraph
    if docx.paragraphs:
        p = docx.paragraphs[0]._element
        p.getparent().remove(p)

    for page_idx, page in enumerate(document.pages):
        logger.info(f"Rendering page {page_idx + 1}/{len(document.pages)}")

        # Create a new section for each page (except the first)
        if page_idx > 0:
            docx.add_section()

        section = docx.sections[-1]

        # Set page dimensions (convert points to EMU)
        section.page_width = int(page.width * PT_TO_EMU)
        section.page_height = int(page.height * PT_TO_EMU)

        # Set zero margins to align page coordinates with PDF points
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)

        # Determine orientation
        if page.width > page.height:
            section.orientation = WD_ORIENT.LANDSCAPE
        else:
            section.orientation = WD_ORIENT.PORTRAIT

        # Single paragraph container per page to keep content on the same page
        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # Add background image if available
        if page.background_image:
            _add_background_image(paragraph, section, page)

        # Add text blocks as positioned textboxes
        if page.blocks:
            _add_text_blocks(paragraph, page, docx)

    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    docx.save(output_path)

    logger.info(f"DOCX saved to: {output_path}")
    return output_path


def _add_background_image(paragraph, section, page: Page):
    """Add the page background image as an inline picture."""
    try:
        # Calculate usable area
        usable_width = (
            section.page_width - section.left_margin - section.right_margin
        )
        usable_height = (
            section.page_height - section.top_margin - section.bottom_margin
        )

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        image_stream = io.BytesIO(page.background_image)

        # Add image stretched to page size
        run.add_picture(image_stream, width=usable_width, height=usable_height)
        _convert_inline_to_anchor(run)

    except Exception as e:
        logger.warning(f"Failed to add background image for page {page.page_number + 1}: {e}")


def _add_text_blocks(paragraph, page: Page, docx: DocxDocument):
    """
    Add text blocks to the document.

    Renders each block as a positioned DrawingML textbox to preserve layout.
    """
    text_blocks = [b for b in page.blocks if b.type == BlockType.TEXT]

    if not text_blocks:
        return

    for block in text_blocks:
        _add_textbox_run(paragraph, block)


def _add_textbox_run(paragraph, block: Block):
    """Insert a DrawingML textbox run positioned at the block's bbox."""
    if not block.content:
        return

    # Ensure minimal size to avoid invalid shapes
    width = max(block.bbox.width, 0.1)
    height = max(block.bbox.height, 0.1)

    x_emu = int(block.bbox.x0 * PT_TO_EMU)
    y_emu = int(block.bbox.y0 * PT_TO_EMU)
    w_emu = int(width * PT_TO_EMU)
    h_emu = int(height * PT_TO_EMU)

    run = paragraph.add_run()
    drawing = OxmlElement("w:drawing")

    anchor = OxmlElement("wp:anchor")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    pos_offset_h = OxmlElement("wp:posOffset")
    pos_offset_h.text = str(x_emu)
    position_h.append(pos_offset_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    pos_offset_v = OxmlElement("wp:posOffset")
    pos_offset_v.text = str(y_emu)
    position_v.append(pos_offset_v)

    extent = OxmlElement("wp:extent")
    extent.set("cx", str(w_emu))
    extent.set("cy", str(h_emu))

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")

    wrap_none = OxmlElement("wp:wrapNone")

    doc_pr = OxmlElement("wp:docPr")
    doc_pr.set("id", str(next(_shape_id_counter)))
    doc_pr.set("name", "TextBox")

    c_nv = OxmlElement("wp:cNvGraphicFramePr")
    locks = OxmlElement("a:graphicFrameLocks")
    locks.set("noChangeAspect", "1")
    c_nv.append(locks)

    graphic = OxmlElement("a:graphic")
    graphic_data = OxmlElement("a:graphicData")
    graphic_data.set("uri", "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")

    wsp = OxmlElement("wps:wsp")
    c_nv_sp_pr = OxmlElement("wps:cNvSpPr")
    c_nv_sp_pr.set("txBox", "1")
    wsp.append(c_nv_sp_pr)

    sp_pr = OxmlElement("wps:spPr")
    xfrm = OxmlElement("a:xfrm")
    off = OxmlElement("a:off")
    off.set("x", "0")
    off.set("y", "0")
    ext = OxmlElement("a:ext")
    ext.set("cx", str(w_emu))
    ext.set("cy", str(h_emu))
    xfrm.append(off)
    xfrm.append(ext)
    sp_pr.append(xfrm)
    prst = OxmlElement("a:prstGeom")
    prst.set("prst", "rect")
    av = OxmlElement("a:avLst")
    prst.append(av)
    sp_pr.append(prst)
    wsp.append(sp_pr)

    txbx = OxmlElement("wps:txbx")
    txbx_content = OxmlElement("w:txbxContent")
    font_name = block.font_name or _select_font_name(block.content)
    font_size = block.font_size
    font_color = block.font_color
    bold = bool(font_size and font_size >= 28)
    lines = block.content.splitlines() if block.content else [""]
    for line in lines:
        txbx_content.append(_build_txbx_paragraph(line, font_name, font_size, font_color, bold))
    txbx.append(txbx_content)
    wsp.append(txbx)

    body_pr = OxmlElement("wps:bodyPr")
    body_pr.set("wrap", "none")
    body_pr.set("lIns", "0")
    body_pr.set("tIns", "0")
    body_pr.set("rIns", "0")
    body_pr.set("bIns", "0")
    body_pr.set("anchor", "t")
    wsp.append(body_pr)

    graphic_data.append(wsp)
    graphic.append(graphic_data)

    anchor.append(simple_pos)
    anchor.append(position_h)
    anchor.append(position_v)
    anchor.append(extent)
    anchor.append(effect_extent)
    anchor.append(wrap_none)
    anchor.append(doc_pr)
    anchor.append(c_nv)
    anchor.append(graphic)

    drawing.append(anchor)
    run._r.append(drawing)


def _convert_inline_to_anchor(run):
    """
    Convert an inline picture in the given run to a floating anchor
    positioned at page origin and behind text.
    """
    try:
        inline = run._r.xpath("./w:drawing/wp:inline")[0]
    except IndexError:
        return

    drawing = inline.getparent()

    anchor = OxmlElement("wp:anchor")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "1")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    pos_offset_h = OxmlElement("wp:posOffset")
    pos_offset_h.text = "0"
    position_h.append(pos_offset_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    pos_offset_v = OxmlElement("wp:posOffset")
    pos_offset_v.text = "0"
    position_v.append(pos_offset_v)

    extent = inline.find(qn("wp:extent"))
    doc_pr = inline.find(qn("wp:docPr"))
    c_nv = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    wrap_none = OxmlElement("wp:wrapNone")

    anchor.append(simple_pos)
    anchor.append(position_h)
    anchor.append(position_v)
    if extent is not None:
        anchor.append(extent)
    if wrap_none is not None:
        anchor.append(wrap_none)
    if doc_pr is not None:
        anchor.append(doc_pr)
    if c_nv is not None:
        anchor.append(c_nv)
    if graphic is not None:
        anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def _build_vml_shapetype() -> OxmlElement:
    """Build the shared VML shapetype definition for textboxes."""
    shapetype = OxmlElement("v:shapetype")
    shapetype.set("id", "_x0000_t202")
    shapetype.set("coordsize", "21600,21600")
    shapetype.set(qn("o:spt"), "202")
    shapetype.set("path", "m,l,21600r21600,l21600,xe")

    stroke = OxmlElement("v:stroke")
    stroke.set("joinstyle", "miter")
    shapetype.append(stroke)

    path = OxmlElement("v:path")
    path.set("gradientshapeok", "t")
    path.set(qn("o:connecttype"), "rect")
    shapetype.append(path)

    return shapetype


def _build_txbx_paragraph(
    text: str,
    font_name: str,
    font_size: Optional[float],
    font_color: Optional[tuple[int, int, int]],
    bold: bool = False,
) -> OxmlElement:
    """Build a paragraph element inside a textbox."""
    p = OxmlElement("w:p")

    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    p_pr.append(spacing)
    p.append(p_pr)

    r_el = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    if font_name:
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_pr.append(r_fonts)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        r_pr.append(sz)

        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), str(int(font_size * 2)))
        r_pr.append(sz_cs)

    if font_color:
        red, green, blue = font_color
        color = OxmlElement("w:color")
        color.set(qn("w:val"), f"{red:02X}{green:02X}{blue:02X}")
        r_pr.append(color)

    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
        bcs = OxmlElement("w:bCs")
        r_pr.append(bcs)

    r_el.append(r_pr)

    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r_el.append(t)
    p.append(r_el)

    return p


def _estimate_alignment(block: Block, page_width: float) -> str:
    """Estimate text alignment from block position."""
    center_x = (block.bbox.x0 + block.bbox.x1) / 2
    page_center = page_width / 2

    # If block is narrow and centered
    block_width = block.bbox.x1 - block.bbox.x0
    if block_width < page_width * 0.6:
        if abs(center_x - page_center) < page_width * 0.1:
            return "center"
        elif block.bbox.x0 > page_width * 0.6:
            return "right"

    return "left"


def _select_font_name(text: str) -> str:
    """Pick a reasonable font based on detected script."""
    if any("\u4e00" <= ch <= "\u9fff" for ch in text):
        return "SimSun"
    return "Arial"
